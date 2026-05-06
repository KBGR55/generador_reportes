"""Generación del Analítico de Estudios (.docx)."""
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportes.config import OUTPUT_DIR
from reportes.core.docx_helpers import set_table_borders, format_date


_SEMESTER_GROUPS = [
    (1, "PRIMER AÑO - I Semestre", "1 ISR, I Sem"),
    (2, "PRIMER AÑO - II Semestre", "1 ISR, II Sem"),
    (3, "SEGUNDO AÑO - I Semestre", "2 ISR, I Sem"),
    (4, "SEGUNDO AÑO - II Semestre", "2 ISR, II Sem"),
    (5, "TERCER AÑO - I Semestre", "3 ISR, I Sem"),
    (6, "TERCER AÑO - II Semestre", "3 ISR, II Sem"),
]

_LATIN_BY_ORDEN = {
    1: ("L 101", "Latín I"),
    2: ("L 102", "Latín II"),
    3: ("L 103", "Latín III"),
    4: ("L 104", "Latín IV"),
    5: ("L 105", "Latín V"),
    6: ("L 106", "Latín VI"),
}

_INSTITUCION = "Estudiantado Santa Mariana de Jesús, Ecuador"
_HEADERS = ["Código", "Materia", "Nota", "Fecha", "Institución"]
_COL_WIDTHS = [Cm(2), Cm(6), Cm(1.5), Cm(2.5), Cm(6)]


def _index_notas(notas, alumna_nombre):
    """Diccionario {materia_normalizada: nota_dict} de la alumna."""
    return {
        n["materia"].strip().lower(): n
        for n in notas
        if n["alumna"] == alumna_nombre
    }


def _materias_del_semestre(orden, plan):
    """Materias del orden N + Latín correspondiente si falta."""
    materias = [m for m in plan if m["orden"] == orden]
    if orden in _LATIN_BY_ORDEN:
        cod, mat = _LATIN_BY_ORDEN[orden]
        if not any(m["codigo"] == cod for m in materias):
            materias.append({
                "codigo": cod, "materia": mat, "orden": orden,
                "anio_desc": "Lingue", "ects": 6,
            })
    return materias


def _add_header_row(table, row_idx, headers, font_size):
    for i, h in enumerate(headers):
        cell = table.rows[row_idx].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)


def _fill_data_row(row, mat, notas_alumna):
    row.cells[0].text = mat["codigo"]
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[1].text = mat["materia"]

    nota_data = notas_alumna.get(mat["materia"].strip().lower())
    if nota_data:
        row.cells[2].text = (str(nota_data["nota"])
                             if nota_data["nota"] is not None else "")
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].text = format_date(nota_data["fecha"])
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[4].text = _INSTITUCION

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)


def _build_semester_table(doc, titulo, materias, notas_alumna):
    table = doc.add_table(rows=len(materias) + 2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Título fusionado
    cell = table.rows[0].cells[0]
    cell.merge(table.rows[0].cells[4])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(9)

    _add_header_row(table, 1, _HEADERS, font_size=8)

    for idx, mat in enumerate(materias):
        _fill_data_row(table.rows[idx + 2], mat, notas_alumna)

    for row in table.rows:
        for i, w in enumerate(_COL_WIDTHS):
            row.cells[i].width = w

    set_table_borders(table)


def _build_extras_table(doc, plan, notas_alumna):
    extras = [m for m in plan if m["orden"] == 20 and m["materia"]]
    if not extras:
        return

    table = doc.add_table(rows=len(extras) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _add_header_row(table, 0, _HEADERS, font_size=8)

    for idx, mat in enumerate(extras):
        row = table.rows[idx + 1]
        row.cells[0].text = mat["codigo"]
        row.cells[1].text = mat["materia"]
        nota_data = notas_alumna.get(mat["materia"].strip().lower())
        if nota_data:
            row.cells[2].text = (str(nota_data["nota"])
                                 if nota_data["nota"] is not None else "")
            row.cells[3].text = format_date(nota_data["fecha"])
        row.cells[4].text = _INSTITUCION
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    set_table_borders(table)


def generate_analitico(alumna_nombre, notas, datos_alumnas, plan):
    """Genera el Analítico y devuelve la ruta del archivo creado."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    dato_alumna = next(
        (d for d in datos_alumnas
         if d["nombre_religioso"] == alumna_nombre or d["alumna"] == alumna_nombre),
        None,
    )

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Certificado de Estudios")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Instituto "Servidoras del Señor y de la Virgen de Matará"')
    run.font.size = Pt(11)

    nombre_mostrar = (
        (dato_alumna.get("nombre_religioso") or
         dato_alumna.get("alumna", alumna_nombre))
        if dato_alumna else alumna_nombre
    )
    p = doc.add_paragraph()
    run = p.add_run("Alumna: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(nombre_mostrar)
    run.font.size = Pt(11)
    run.underline = True

    notas_alumna = _index_notas(notas, alumna_nombre)

    for orden, titulo, _anio_desc in _SEMESTER_GROUPS:
        materias_sem = _materias_del_semestre(orden, plan)
        if not materias_sem:
            continue
        doc.add_paragraph("")
        _build_semester_table(doc, titulo, materias_sem, notas_alumna)

    # Cursos extra
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("CURSOS EXTRA:")
    run.bold = True
    run.font.size = Pt(10)

    _build_extras_table(doc, plan, notas_alumna)

    # Firmas
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("________________\t\t\t\t\t___________________")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Superiora Local\t\t\t\t\tEncargada de Estudios")
    run.font.size = Pt(10)

    safe_name = alumna_nombre.replace(" ", "_")[:30]
    filename = f"Analitico_{safe_name}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
