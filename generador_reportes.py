#!/usr/bin/env python3
"""
Generador de Reportes - Estudiantado Santa Mariana de Jesús
Genera Actas de Examen y Analíticos de Estudios desde esquema.xlsx
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import openpyxl
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime
import subprocess
import sys

EXCEL_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "esquema.xlsx")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reportes_generados")


def load_excel_data():
    """Load all data from the Excel file."""
    wb = openpyxl.load_workbook(EXCEL_PATH, data_only=True)

    # Notas
    ws = wb["notas"]
    notas = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None and row[1] is None:
            continue
        notas.append({
            "profesor": row[0] or "",
            "alumna": row[1] or "",
            "materia": row[2] or "",
            "anio": row[3],
            "semestre": row[4] or "",
            "fecha": row[5],
            "nota": row[6],
        })

    # Datos de alumnas
    ws = wb["datos"]
    datos = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[2] is None and row[1] is None:
            continue
        datos.append({
            "nacionalidad": row[0] or "",
            "nombre_religioso": row[1] or "",
            "alumna": row[2] or "",
            "num_documento": row[3] or "",
            "tipo_documento": row[4] or "",
            "fecha_nacimiento": row[5],
        })

    # Profesores
    ws = wb["profesores"]
    profesores = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[1]:
            profesores.append(row[1])

    # Plan de estudios
    ws = wb["Plan de estudios"]
    plan = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[1] is None:
            continue
        codigo = (row[0] or "").strip()
        plan.append({
            "codigo": codigo,
            "materia": row[1],
            "orden": row[2],
            "anio_desc": row[3] or "",
            "ects": row[4],
        })

    wb.close()
    return notas, datos, profesores, plan


def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz": 4, "val": "single"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """Apply borders to entire table."""
    border = {"sz": 4, "val": "single", "color": "000000", "space": "0"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def format_date(d):
    """Format a date or datetime to DD/MM/YYYY."""
    if isinstance(d, datetime.datetime):
        return d.strftime("%d/%m/%Y")
    elif isinstance(d, datetime.date):
        return d.strftime("%d/%m/%Y")
    elif isinstance(d, str):
        return d
    return ""


def anio_academico_from_nota(anio, semestre, plan):
    """Determine academic year label (Primero, Segundo, Tercero) from year number."""
    mapping = {2023: "Primero", 2024: "Segundo", 2025: "Tercero"}
    return mapping.get(anio, str(anio) if anio else "")


def generate_acta_examen(materia, profesor, anio, semestre, fecha, notas_filtradas, datos_alumnas):
    """Generate Acta de Examen document."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Map year to academic year name
    anio_map = {"2023": "Primero", "2024": "Segundo", "2025": "Tercero"}
    anio_academico = anio_map.get(str(anio), str(anio))

    # === Header section (NO borders) ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Estudiantado Santa Mariana de Jesús")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")

    # Info fields as simple paragraphs (no table borders)
    info_fields = [
        ("Año:", str(anio)),
        ("Materia:", materia),
        ("Profesor:", profesor),
        ("Fecha:", format_date(fecha)),
        ("Año académico:", anio_academico),
        ("Semestre:", str(semestre)),
    ]
    for label, value in info_fields:
        p = doc.add_paragraph()
        run = p.add_run(label + "  ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.size = Pt(11)

    doc.add_paragraph("")

    # === Student table (WITH borders) ===
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row for students
    row = table.add_row()
    headers = ["N°", "Apellido, nombre", "Documento", "Nota"]
    for i, h in enumerate(headers):
        p = row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        pass  # No background color

    # Build lookup for alumna data
    datos_map = {}
    for d in datos_alumnas:
        datos_map[d["nombre_religioso"]] = d
        datos_map[d["alumna"]] = d

    # Student rows (at least 12 rows like template)
    num_rows = max(12, len(notas_filtradas))
    for i in range(num_rows):
        row = table.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i < len(notas_filtradas):
            nota = notas_filtradas[i]
            alumna_name = nota["alumna"]
            # Get nombre religioso (nombre de monja)
            dato = datos_map.get(alumna_name, {})
            nombre_mostrar = dato.get("nombre_religioso", "") or alumna_name
            doc_num = dato.get("num_documento", "")
            row.cells[1].text = nombre_mostrar
            row.cells[2].text = str(doc_num)
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[3].text = str(nota["nota"]) if nota["nota"] is not None else ""
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set font size for student table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is None:
                        run.font.size = Pt(10)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(4)
        row.cells[3].width = Cm(2)

    set_table_borders(table)

    # Signature line
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(".............................................")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Firma del titular")
    run.font.size = Pt(10)

    # Observations
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Observaciones:")
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph("_" * 80)
    doc.add_paragraph("_" * 80)

    # Save
    safe_materia = materia.replace("/", "-").replace(":", "-")[:40]
    filename = f"Acta_Examen_{safe_materia}_{anio}_{semestre}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


def generate_analitico(alumna_nombre, notas, datos_alumnas, plan):
    """Generate Analitico de Estudios document."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Get student data
    dato_alumna = None
    for d in datos_alumnas:
        if d["nombre_religioso"] == alumna_nombre or d["alumna"] == alumna_nombre:
            dato_alumna = d
            break

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Certificado de Estudios")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Instituto "Servidoras del Señor y de la Virgen de Matará"')
    run.font.size = Pt(11)

    # Student name - mostrar nombre religioso
    nombre_mostrar = (dato_alumna.get("nombre_religioso") or dato_alumna.get("alumna", alumna_nombre)) if dato_alumna else alumna_nombre
    p = doc.add_paragraph()
    run = p.add_run(f"Alumna: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(nombre_mostrar)
    run.font.size = Pt(11)
    run.underline = True

    # Build notas lookup by materia name (normalize for matching)
    notas_alumna = {}
    for n in notas:
        if n["alumna"] == alumna_nombre:
            notas_alumna[n["materia"].strip().lower()] = n

    # Define semester groups matching the template structure
    semester_groups = [
        (1, "PRIMER AÑO - I Semestre", "1 ISR, I Sem"),
        (2, "PRIMER AÑO - II Semestre", "1 ISR, II Sem"),
        (3, "SEGUNDO AÑO - I Semestre", "2 ISR, I Sem"),
        (4, "SEGUNDO AÑO - II Semestre", "2 ISR, II Sem"),
        (5, "TERCER AÑO - I Semestre", "3 ISR, I Sem"),
        (6, "TERCER AÑO - II Semestre", "3 ISR, II Sem"),
    ]

    institucion = "Estudiantado Santa Mariana de Jesús, Ecuador"

    for orden, titulo, anio_desc in semester_groups:
        # Filter plan entries for this semester
        materias_sem = [m for m in plan if m["orden"] == orden]
        # Add Latin for relevant semesters
        latin_map = {1: ("L 101", "Latín I"), 2: ("L 102", "Latín II"),
                     3: ("L 103", "Latín III"), 4: ("L 104", "Latín IV"),
                     5: ("L 105", "Latín V"), 6: ("L 106", "Latín VI")}
        if orden in latin_map:
            cod, mat = latin_map[orden]
            # Check it's not already there
            if not any(m["codigo"] == cod for m in materias_sem):
                materias_sem.append({"codigo": cod, "materia": mat, "orden": orden,
                                     "anio_desc": "Lingue", "ects": 6})

        if not materias_sem:
            continue

        doc.add_paragraph("")

        # Create table
        num_rows = len(materias_sem) + 2  # title + header + data
        table = doc.add_table(rows=num_rows, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Title row
        cell = table.rows[0].cells[0]
        cell.merge(table.rows[0].cells[4])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(9)

        # Header row
        headers = ["Código", "Materia", "Nota", "Fecha", "Institución"]
        for i, h in enumerate(headers):
            cell = table.rows[1].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(8)
            pass  # No background color

        # Data rows
        for idx, mat in enumerate(materias_sem):
            row = table.rows[idx + 2]
            row.cells[0].text = mat["codigo"]
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = mat["materia"]

            # Look up grade
            mat_key = mat["materia"].strip().lower()
            nota_data = notas_alumna.get(mat_key)
            if nota_data:
                row.cells[2].text = str(nota_data["nota"]) if nota_data["nota"] is not None else ""
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[3].text = format_date(nota_data["fecha"])
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[4].text = institucion

            # Font size
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        # Set column widths
        widths = [Cm(2), Cm(6), Cm(1.5), Cm(2.5), Cm(6)]
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

        set_table_borders(table)

    # Extra courses section
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("CURSOS EXTRA:")
    run.bold = True
    run.font.size = Pt(10)

    # Extra materias (orden 20)
    extras = [m for m in plan if m["orden"] == 20 and m["materia"]]
    if extras:
        table = doc.add_table(rows=len(extras) + 1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Código", "Materia", "Nota", "Fecha", "Institución"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(8)
            pass  # No background color

        for idx, mat in enumerate(extras):
            row = table.rows[idx + 1]
            row.cells[0].text = mat["codigo"]
            row.cells[1].text = mat["materia"]
            mat_key = mat["materia"].strip().lower()
            nota_data = notas_alumna.get(mat_key)
            if nota_data:
                row.cells[2].text = str(nota_data["nota"]) if nota_data["nota"] is not None else ""
                row.cells[3].text = format_date(nota_data["fecha"])
            row.cells[4].text = institucion
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        set_table_borders(table)

    # Signatures
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("________________\t\t\t\t\t___________________")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Superiora Local\t\t\t\t\tEncargada de Estudios")
    run.font.size = Pt(10)

    # Save
    safe_name = alumna_nombre.replace(" ", "_")[:30]
    filename = f"Analitico_{safe_name}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


class ReportGeneratorApp:
    # Color palette (Tailwind-inspired)
    BG = "#f1f5f9"        # slate-100
    CARD = "#ffffff"
    PRIMARY = "#4f46e5"   # indigo-600
    PRIMARY_DARK = "#4338ca"
    TEXT = "#0f172a"      # slate-900
    MUTED = "#64748b"     # slate-500
    BORDER = "#e2e8f0"    # slate-200
    SUCCESS = "#059669"   # emerald-600
    ACCENT_BAR = "#6366f1"  # indigo-500

    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Reportes - Estudiantado Santa Mariana de Jesús")
        self.root.geometry("760x640")
        self.root.minsize(720, 600)
        self.root.configure(bg=self.BG)

        # Load data
        try:
            self.notas, self.datos, self.profesores, self.plan = load_excel_data()
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo leer esquema.xlsx:\n{e}")
            sys.exit(1)

        # Get unique values for dropdowns
        self.materias_unicas = sorted(set(n["materia"] for n in self.notas if n["materia"]))
        self.alumnas_unicas = sorted(set(n["alumna"] for n in self.notas if n["alumna"]))
        self.profesores_unicos = sorted(set(n["profesor"] for n in self.notas if n["profesor"]))
        # Also add from datos sheet
        for d in self.datos:
            if d["nombre_religioso"] and d["nombre_religioso"] not in self.alumnas_unicas:
                self.alumnas_unicas.append(d["nombre_religioso"])
        # Add all materias from plan too
        all_materias = sorted(set(
            list(self.materias_unicas) +
            [m["materia"] for m in self.plan if m["materia"]]
        ))
        self.materias_unicas = all_materias

        self.build_ui()

    def build_ui(self):
        BG, CARD, PRIMARY, PRIMARY_DARK = self.BG, self.CARD, self.PRIMARY, self.PRIMARY_DARK
        TEXT, MUTED, BORDER, SUCCESS = self.TEXT, self.MUTED, self.BORDER, self.SUCCESS

        style = ttk.Style()
        style.theme_use("clam")

        # Base
        style.configure(".", background=BG, foreground=TEXT, font=("Helvetica", 10))
        style.configure("TFrame", background=BG)
        style.configure("TLabel", background=BG, foreground=TEXT)

        # Header text
        style.configure("Title.TLabel", font=("Helvetica", 19, "bold"),
                        background=BG, foreground=TEXT)
        style.configure("Subtitle.TLabel", font=("Helvetica", 10),
                        background=BG, foreground=MUTED)
        style.configure("Info.TLabel", font=("Helvetica", 9),
                        background=BG, foreground=MUTED)

        # Card-context labels
        style.configure("Section.TLabel", font=("Helvetica", 12, "bold"),
                        background=CARD, foreground=TEXT)
        style.configure("Form.TLabel", font=("Helvetica", 10),
                        background=CARD, foreground=TEXT)
        style.configure("Card.TLabel", background=CARD, foreground=TEXT)

        # Input fields
        style.configure("TEntry", fieldbackground=CARD, bordercolor=BORDER,
                        lightcolor=BORDER, darkcolor=BORDER, padding=6)
        style.configure("TCombobox", fieldbackground=CARD, bordercolor=BORDER,
                        lightcolor=BORDER, darkcolor=BORDER, padding=4,
                        arrowsize=14)
        style.map("TCombobox",
                  fieldbackground=[("readonly", CARD), ("disabled", "#f8fafc")],
                  bordercolor=[("focus", PRIMARY)])
        style.map("TEntry", bordercolor=[("focus", PRIMARY)])

        # Radiobutton on card background
        style.configure("Type.TRadiobutton", background=CARD, foreground=TEXT,
                        font=("Helvetica", 11), padding=4)
        style.map("Type.TRadiobutton",
                  background=[("active", CARD)],
                  foreground=[("selected", PRIMARY)])

        # Primary action button
        style.configure("Primary.TButton",
                        background=PRIMARY, foreground="white",
                        font=("Helvetica", 12, "bold"),
                        padding=(28, 12), borderwidth=0, relief="flat",
                        focuscolor=PRIMARY)
        style.map("Primary.TButton",
                  background=[("active", PRIMARY_DARK), ("pressed", PRIMARY_DARK),
                              ("disabled", "#94a3b8")],
                  foreground=[("disabled", "white")])

        # Secondary outline button
        style.configure("Secondary.TButton",
                        background=CARD, foreground=PRIMARY,
                        font=("Helvetica", 10, "bold"),
                        padding=(14, 8), borderwidth=1, relief="flat",
                        bordercolor=BORDER, focuscolor=CARD)
        style.map("Secondary.TButton",
                  background=[("active", "#eef2ff")],
                  bordercolor=[("active", PRIMARY)])

        # Treeview (lista de reportes)
        style.configure("Treeview", background=CARD, fieldbackground=CARD,
                        foreground=TEXT, rowheight=26, borderwidth=0,
                        font=("Helvetica", 10))
        style.configure("Treeview.Heading", background="#f8fafc",
                        foreground=TEXT, font=("Helvetica", 10, "bold"),
                        relief="flat", padding=6)
        style.map("Treeview", background=[("selected", PRIMARY)],
                  foreground=[("selected", "white")])

        # Status label (success)
        style.configure("Status.TLabel", background=BG, foreground=SUCCESS,
                        font=("Helvetica", 10, "bold"))

        # === Header (with thin accent bar) ===
        accent = tk.Frame(self.root, bg=self.ACCENT_BAR, height=4)
        accent.pack(fill="x", side="top")

        header = ttk.Frame(self.root)
        header.pack(fill="x", padx=28, pady=(18, 4))

        header_left = ttk.Frame(header)
        header_left.pack(side="left", anchor="w")
        ttk.Label(header_left, text="Generador de Reportes",
                  style="Title.TLabel").pack(anchor="w")
        ttk.Label(header_left, text="Estudiantado Santa Mariana de Jesús",
                  style="Subtitle.TLabel").pack(anchor="w", pady=(2, 0))

        ttk.Button(header, text="Ver reportes generados",
                   style="Secondary.TButton", cursor="hand2",
                   command=self.open_reports_viewer).pack(side="right", anchor="ne")

        # Divider
        tk.Frame(self.root, height=1, bg=BORDER).pack(fill="x", padx=28, pady=(12, 16))

        # === Type selector card ===
        type_card = tk.Frame(self.root, bg=CARD,
                             highlightbackground=BORDER, highlightthickness=1, bd=0)
        type_card.pack(padx=28, pady=(0, 14), fill="x")

        type_inner = tk.Frame(type_card, bg=CARD)
        type_inner.pack(padx=20, pady=16, fill="x")

        ttk.Label(type_inner, text="Tipo de reporte",
                  style="Section.TLabel").pack(anchor="w", pady=(0, 10))

        self.report_type = tk.StringVar(value="acta")
        radio_row = tk.Frame(type_inner, bg=CARD)
        radio_row.pack(anchor="w")
        ttk.Radiobutton(radio_row, text="Acta de Examen",
                        variable=self.report_type, value="acta",
                        style="Type.TRadiobutton",
                        command=self.toggle_panels).pack(side="left", padx=(0, 36))
        ttk.Radiobutton(radio_row, text="Analítico de Estudios",
                        variable=self.report_type, value="analitico",
                        style="Type.TRadiobutton",
                        command=self.toggle_panels).pack(side="left")

        # === ACTA card ===
        self.acta_frame = tk.Frame(self.root, bg=CARD,
                                   highlightbackground=BORDER, highlightthickness=1, bd=0)
        acta_inner = tk.Frame(self.acta_frame, bg=CARD)
        acta_inner.pack(padx=20, pady=16, fill="x", expand=True)

        ttk.Label(acta_inner, text="Datos del Acta de Examen",
                  style="Section.TLabel").pack(anchor="w", pady=(0, 14))

        form_acta = tk.Frame(acta_inner, bg=CARD)
        form_acta.pack(fill="x")
        form_acta.columnconfigure(1, weight=1)

        row = 0
        ttk.Label(form_acta, text="Materia",
                  style="Form.TLabel").grid(row=row, column=0, sticky="w",
                                            pady=7, padx=(0, 14))
        self.acta_materia = ttk.Combobox(form_acta, values=self.materias_unicas)
        self.acta_materia.grid(row=row, column=1, sticky="ew", pady=7)

        row += 1
        ttk.Label(form_acta, text="Profesor",
                  style="Form.TLabel").grid(row=row, column=0, sticky="w",
                                            pady=7, padx=(0, 14))
        self.acta_profesor = ttk.Combobox(form_acta,
                                          values=self.profesores_unicos + self.profesores)
        self.acta_profesor.grid(row=row, column=1, sticky="ew", pady=7)

        row += 1
        ttk.Label(form_acta, text="Año",
                  style="Form.TLabel").grid(row=row, column=0, sticky="w",
                                            pady=7, padx=(0, 14))
        self.acta_anio = ttk.Combobox(form_acta,
                                      values=["2023", "2024", "2025", "2026"], width=14)
        self.acta_anio.grid(row=row, column=1, sticky="w", pady=7)
        self.acta_anio.set("2025")

        row += 1
        ttk.Label(form_acta, text="Semestre",
                  style="Form.TLabel").grid(row=row, column=0, sticky="w",
                                            pady=7, padx=(0, 14))
        self.acta_semestre = ttk.Combobox(form_acta, values=["I", "II"], width=14)
        self.acta_semestre.grid(row=row, column=1, sticky="w", pady=7)
        self.acta_semestre.set("I")

        row += 1
        ttk.Label(form_acta, text="Fecha (DD/MM/AAAA)",
                  style="Form.TLabel").grid(row=row, column=0, sticky="w",
                                            pady=7, padx=(0, 14))
        self.acta_fecha = ttk.Entry(form_acta, width=22)
        self.acta_fecha.grid(row=row, column=1, sticky="w", pady=7)
        today = datetime.date.today().strftime("%d/%m/%Y")
        self.acta_fecha.insert(0, today)

        # === ANALITICO card ===
        self.analitico_frame = tk.Frame(self.root, bg=CARD,
                                        highlightbackground=BORDER, highlightthickness=1, bd=0)
        ana_inner = tk.Frame(self.analitico_frame, bg=CARD)
        ana_inner.pack(padx=20, pady=16, fill="x", expand=True)

        ttk.Label(ana_inner, text="Datos del Analítico de Estudios",
                  style="Section.TLabel").pack(anchor="w", pady=(0, 14))

        form_ana = tk.Frame(ana_inner, bg=CARD)
        form_ana.pack(fill="x")
        form_ana.columnconfigure(1, weight=1)

        ttk.Label(form_ana, text="Alumna",
                  style="Form.TLabel").grid(row=0, column=0, sticky="w",
                                            pady=7, padx=(0, 14))
        self.analitico_alumna = ttk.Combobox(form_ana, values=self.alumnas_unicas)
        self.analitico_alumna.grid(row=0, column=1, sticky="ew", pady=7)
        if self.alumnas_unicas:
            self.analitico_alumna.set(self.alumnas_unicas[0])

        # Status label
        self.status_var = tk.StringVar(value="")
        self.status_label = ttk.Label(self.root, textvariable=self.status_var,
                                      style="Status.TLabel")

        # Generate button
        self.generate_btn = ttk.Button(self.root, text="Generar Reporte",
                                       style="Primary.TButton",
                                       command=self.generate_report,
                                       cursor="hand2")

        # Info bar at bottom
        self.info_var = tk.StringVar(
            value=f"Datos cargados: {len(self.notas)} notas  ·  "
                  f"{len(self.datos)} alumnas  ·  {len(self.plan)} materias en plan"
        )
        info_bar = ttk.Frame(self.root)
        info_bar.pack(side="bottom", fill="x", padx=28, pady=(0, 14))
        tk.Frame(info_bar, height=1, bg=BORDER).pack(fill="x", pady=(0, 8))
        ttk.Label(info_bar, textvariable=self.info_var,
                  style="Info.TLabel").pack(anchor="w")

        # Initial layout
        self.toggle_panels()

    def toggle_panels(self):
        self.acta_frame.pack_forget()
        self.analitico_frame.pack_forget()
        self.status_label.pack_forget()
        self.generate_btn.pack_forget()

        if self.report_type.get() == "acta":
            self.acta_frame.pack(padx=28, pady=(0, 14), fill="x")
        else:
            self.analitico_frame.pack(padx=28, pady=(0, 14), fill="x")

        self.status_label.pack(padx=28, pady=(4, 0))
        self.generate_btn.pack(pady=(14, 8))

    def open_reports_viewer(self):
        win = tk.Toplevel(self.root)
        win.title("Reportes generados")
        win.geometry("760x480")
        win.minsize(700, 420)
        win.configure(bg=self.BG)
        win.transient(self.root)

        tk.Frame(win, bg=self.ACCENT_BAR, height=4).pack(fill="x", side="top")

        head = ttk.Frame(win)
        head.pack(fill="x", padx=24, pady=(16, 6))
        ttk.Label(head, text="Reportes generados",
                  style="Title.TLabel").pack(anchor="w")
        ttk.Label(head, text=OUTPUT_DIR,
                  style="Subtitle.TLabel").pack(anchor="w", pady=(2, 0))

        tk.Frame(win, height=1, bg=self.BORDER).pack(fill="x", padx=24, pady=(10, 12))

        card = tk.Frame(win, bg=self.CARD,
                        highlightbackground=self.BORDER, highlightthickness=1, bd=0)
        card.pack(padx=24, pady=(0, 12), fill="both", expand=True)

        list_wrap = tk.Frame(card, bg=self.CARD)
        list_wrap.pack(padx=12, pady=12, fill="both", expand=True)

        cols = ("nombre", "tipo", "fecha", "tamano")
        tree = ttk.Treeview(list_wrap, columns=cols, show="headings", height=10)
        tree.heading("nombre", text="Nombre")
        tree.heading("tipo", text="Tipo")
        tree.heading("fecha", text="Modificado")
        tree.heading("tamano", text="Tamaño")
        tree.column("nombre", width=360, anchor="w")
        tree.column("tipo", width=110, anchor="w")
        tree.column("fecha", width=140, anchor="w")
        tree.column("tamano", width=80, anchor="e")

        sb = ttk.Scrollbar(list_wrap, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=sb.set)
        tree.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        empty_var = tk.StringVar(value="")
        empty_lbl = ttk.Label(card, textvariable=empty_var,
                              background=self.CARD, foreground=self.MUTED,
                              font=("Helvetica", 10, "italic"))

        def populate():
            for item in tree.get_children():
                tree.delete(item)
            empty_lbl.pack_forget()

            if not os.path.isdir(OUTPUT_DIR):
                empty_var.set("Aún no se han generado reportes.")
                empty_lbl.pack(pady=(0, 14))
                return

            archivos = [f for f in os.listdir(OUTPUT_DIR)
                        if os.path.isfile(os.path.join(OUTPUT_DIR, f))]
            archivos.sort(key=lambda f: os.path.getmtime(
                os.path.join(OUTPUT_DIR, f)), reverse=True)

            if not archivos:
                empty_var.set("Aún no se han generado reportes.")
                empty_lbl.pack(pady=(0, 14))
                return

            for fname in archivos:
                fpath = os.path.join(OUTPUT_DIR, fname)
                if fname.startswith("Acta_"):
                    tipo = "Acta de Examen"
                elif fname.startswith("Analitico"):
                    tipo = "Analítico"
                else:
                    tipo = "Otro"
                mtime = datetime.datetime.fromtimestamp(os.path.getmtime(fpath))
                size_kb = os.path.getsize(fpath) / 1024
                tree.insert("", "end", values=(
                    fname, tipo,
                    mtime.strftime("%d/%m/%Y %H:%M"),
                    f"{size_kb:.0f} KB"))

        def open_path(path):
            try:
                if sys.platform == "linux":
                    subprocess.Popen(["xdg-open", path])
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", path])
                elif sys.platform == "win32":
                    os.startfile(path)
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo abrir:\n{e}")

        def open_selected(_event=None):
            sel = tree.selection()
            if not sel:
                return
            fname = tree.item(sel[0], "values")[0]
            open_path(os.path.join(OUTPUT_DIR, fname))

        def open_folder():
            os.makedirs(OUTPUT_DIR, exist_ok=True)
            open_path(OUTPUT_DIR)

        def delete_selected():
            sel = tree.selection()
            if not sel:
                return
            fname = tree.item(sel[0], "values")[0]
            if not messagebox.askyesno(
                    "Eliminar reporte",
                    f"¿Eliminar definitivamente «{fname}»?", parent=win):
                return
            try:
                os.remove(os.path.join(OUTPUT_DIR, fname))
                populate()
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo eliminar:\n{e}")

        tree.bind("<Double-1>", open_selected)
        tree.bind("<Return>", open_selected)

        btns = tk.Frame(win, bg=self.BG)
        btns.pack(fill="x", padx=24, pady=(0, 18))
        ttk.Button(btns, text="Abrir", style="Primary.TButton",
                   cursor="hand2", command=open_selected).pack(side="left")
        ttk.Button(btns, text="Abrir carpeta", style="Secondary.TButton",
                   cursor="hand2", command=open_folder).pack(side="left", padx=(10, 0))
        ttk.Button(btns, text="Refrescar", style="Secondary.TButton",
                   cursor="hand2", command=populate).pack(side="left", padx=(10, 0))
        ttk.Button(btns, text="Eliminar", style="Secondary.TButton",
                   cursor="hand2", command=delete_selected).pack(side="left", padx=(10, 0))
        ttk.Button(btns, text="Cerrar", style="Secondary.TButton",
                   cursor="hand2", command=win.destroy).pack(side="right")

        populate()

    def generate_report(self):
        try:
            # Reload data each time in case Excel was updated
            self.notas, self.datos, self.profesores, self.plan = load_excel_data()

            if self.report_type.get() == "acta":
                filepath = self.generate_acta()
            else:
                filepath = self.generate_analitico()

            self.status_var.set(f"Reporte generado: {os.path.basename(filepath)}")
            messagebox.showinfo("Éxito",
                                f"Reporte generado exitosamente!\n\n{filepath}")

            # Open the file
            if sys.platform == "linux":
                subprocess.Popen(["xdg-open", filepath])
            elif sys.platform == "darwin":
                subprocess.Popen(["open", filepath])
            elif sys.platform == "win32":
                os.startfile(filepath)

        except Exception as e:
            messagebox.showerror("Error", f"Error al generar reporte:\n{e}")
            import traceback
            traceback.print_exc()

    def generate_acta(self):
        materia = self.acta_materia.get().strip()
        profesor = self.acta_profesor.get().strip()
        anio = self.acta_anio.get().strip()
        semestre = self.acta_semestre.get().strip()
        fecha_str = self.acta_fecha.get().strip()

        if not materia:
            raise ValueError("Seleccione una materia")
        if not profesor:
            raise ValueError("Seleccione un profesor")

        # Parse fecha
        try:
            parts = fecha_str.split("/")
            fecha = datetime.date(int(parts[2]), int(parts[1]), int(parts[0]))
        except Exception:
            fecha = fecha_str

        # Filter notas for this materia
        notas_filtradas = []
        for n in self.notas:
            materia_match = n["materia"].strip().lower() == materia.strip().lower()
            if materia_match:
                notas_filtradas.append(n)

        return generate_acta_examen(materia, profesor, anio, semestre, fecha, notas_filtradas, self.datos)

    def generate_analitico(self):
        alumna = self.analitico_alumna.get().strip()
        if not alumna:
            raise ValueError("Seleccione una alumna")

        return generate_analitico(alumna, self.notas, self.datos, self.plan)


if __name__ == "__main__":
    root = tk.Tk()
    app = ReportGeneratorApp(root)
    root.mainloop()