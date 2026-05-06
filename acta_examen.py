"""Generación del Acta de Examen (.docx)."""
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import OUTPUT_DIR
from docx_helpers import set_table_borders, format_date


_ANIO_ACADEMICO = {"2023": "Primero", "2024": "Segundo", "2025": "Tercero"}


def generate_acta_examen(materia, profesor, anio, semestre, fecha,
                         notas_filtradas, datos_alumnas):
    """Genera el Acta de Examen y devuelve la ruta del archivo creado."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    anio_academico = _ANIO_ACADEMICO.get(str(anio), str(anio))

    # Encabezado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Estudiantado Santa Mariana de Jesús")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")

    info_fields = [
        ("Año:", str(anio)),
        ("Materia:", materia),
        ("Profesor:", profesor),
        ("Fecha:", format_date(fecha)),
        ("Año académico:", anio_academico),
        ("Semestre:", str(semestre)),
    ]
    for label, value in info_fields:
        p = doc.add_paragraph()
        run = p.add_run(label + "  ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.size = Pt(11)

    doc.add_paragraph("")

    # Tabla de alumnas
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row = table.add_row()
    headers = ["N°", "Apellido, nombre", "Documento", "Nota"]
    for i, h in enumerate(headers):
        p = row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    datos_map = {}
    for d in datos_alumnas:
        datos_map[d["nombre_religioso"]] = d
        datos_map[d["alumna"]] = d

    num_rows = max(12, len(notas_filtradas))
    for i in range(num_rows):
        row = table.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i < len(notas_filtradas):
            nota = notas_filtradas[i]
            alumna_name = nota["alumna"]
            dato = datos_map.get(alumna_name, {})
            nombre_mostrar = dato.get("nombre_religioso", "") or alumna_name
            doc_num = dato.get("num_documento", "")
            row.cells[1].text = nombre_mostrar
            row.cells[2].text = str(doc_num)
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[3].text = (str(nota["nota"])
                                 if nota["nota"] is not None else "")
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is None:
                        run.font.size = Pt(10)

    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(4)
        row.cells[3].width = Cm(2)

    set_table_borders(table)

    # Firma
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("." * 45)
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Firma del titular")
    run.font.size = Pt(10)

    # Observaciones
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Observaciones:")
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph("_" * 80)
    doc.add_paragraph("_" * 80)

    safe_materia = materia.replace("/", "-").replace(":", "-")[:40]
    filename = f"Acta_Examen_{safe_materia}_{anio}_{semestre}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
